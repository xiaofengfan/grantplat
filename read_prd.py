# -*- coding: utf-8 -*-
from docx import Document

doc = Document(r'e:\DEVWORKSPACE\gant\grantplat\prd\新一代股票量化交易平台功能详细设计 - 副本.docx')

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

for table in doc.tables:
    print("\n=== TABLE ===")
    for row in table.rows:
        row_text = [cell.text for cell in row.cells]
        print(" | ".join(row_text))